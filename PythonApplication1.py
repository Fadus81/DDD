from email import message
from email.contentmanager import ContentManager
from tkinter import Button
import docx
import os
import datetime
import random
import re
import sys
if sys.platform.startswith('win'):
    from docx2pdf import convert
    def file_convert_docx_pdf(dirs):
        file_in_dir = os.listdir(dirs)

   # if not os.path.isdir(f'{dirs}/convert_pdf'):
     #   os.mkdir(f'{dirs}/convert_pdf')

        for file in file_in_dir:
            if file.endswith('.docx'):
                file_k = f'{file.split(".")[0].replace(".", "_")}.pdf'
                convert(f'{dirs}/{file}', f'{dirs}/{file_k}')
                os.remove(f'{dirs}/'+file)
            else:
                continue
elif sys.platform.startswith('linux'):
    import subprocess
    def file_convert_docx_pdf(dirs):
        file_in_dir = os.listdir(dirs)

   # if not os.path.isdir(f'{dirs}/convert_pdf'):
     #   os.mkdir(f'{dirs}/convert_pdf')

        for file in file_in_dir:
            if file.endswith('.docx'):
                file_k =dirs+ f'/{file.split(".")[0].replace(".", "_")}.pdf'
                file_docx = f'{dirs}/{file}'
                try:
                    subprocess.run(["unoconv", "-f", "pdf", "-o", file_k, file_docx], check=True)
                except subprocess.CalledProcessError as err:
                    print(f"Error converting {dirs}/{file} to PDF: {err}")
                os.remove(f'{dirs}/'+file)
            else:
                continue

from aiogram import Bot, types
from aiogram.dispatcher import Dispatcher
from aiogram.dispatcher import FSMContext
from aiogram.utils import executor
from aiogram.dispatcher.filters import Command,Text
from aiogram.types import Message,ChatPermissions
from aiogram.dispatcher.filters.state import State, StatesGroup
from aiogram.contrib.fsm_storage.memory import MemoryStorage
from aiogram.types import InputFile, User
from aiogram.utils.exceptions import ChatNotFound
from aiogram.types import ReplyKeyboardRemove, \
                          ReplyKeyboardMarkup, KeyboardButton, \
                          InlineKeyboardMarkup, InlineKeyboardButton

#from loader import dp

file1 = open("word1/token.txt", "r")
lines = file1.readline()
bot = Bot(token=lines)
file1.close
dp = Dispatcher(bot, storage=MemoryStorage())


def docx_replace(doc, data):
    paragraphs = list(doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraphs.append(paragraph)
    for p in paragraphs:
        for key, val in data.items():
            key_name = '${{{}}}'.format(key) # I'm using placeholders in the form ${PlaceholderName}
            if key_name in p.text:
                inline = p.runs
                # Replace strings and retain the same style.
                # The text to be replaced can be split over several runs so
                # search through, identify which runs need to have text replaced
                # then replace the text in those identified
                started = False
                key_index = 0
                # found_runs is a list of (inline index, index of match, length of match)
                found_runs = list()
                found_all = False
                replace_done = False
                for i in range(len(inline)):

                    # case 1: found in single run so short circuit the replace
                    if key_name in inline[i].text and not started:
                        found_runs.append((i, inline[i].text.find(key_name), len(key_name)))
                        text = inline[i].text.replace(key_name, str(val))
                        inline[i].text = text
                        replace_done = True
                        found_all = True
                        break

                    if key_name[key_index] not in inline[i].text and not started:
                        # keep looking ...
                        continue

                    # case 2: search for partial text, find first run
                    if key_name[key_index] in inline[i].text and inline[i].text[-1] in key_name and not started:
                        # check sequence
                        start_index = inline[i].text.find(key_name[key_index])
                        check_length = len(inline[i].text)
                        for text_index in range(start_index, check_length):
                            if inline[i].text[text_index] != key_name[key_index]:
                                # no match so must be false positive
                                break
                        if key_index == 0:
                            started = True
                        chars_found = check_length - start_index
                        key_index += chars_found
                        found_runs.append((i, start_index, chars_found))
                        if key_index != len(key_name):
                            continue
                        else:
                            # found all chars in key_name
                            found_all = True
                            break

                    # case 2: search for partial text, find subsequent run
                    if key_name[key_index] in inline[i].text and started and not found_all:
                        # check sequence
                        chars_found = 0
                        check_length = len(inline[i].text)
                        for text_index in range(0, check_length):
                            if inline[i].text[text_index] == key_name[key_index]:
                                key_index += 1
                                chars_found += 1
                            else:
                                break
                        # no match so must be end
                        found_runs.append((i, 0, chars_found))
                        if key_index == len(key_name):
                            found_all = True
                            break

                if found_all and not replace_done:
                    for i, item in enumerate(found_runs):
                        index, start, length = [t for t in item]
                        if i == 0:
                            text = inline[index].text.replace(inline[index].text[start:start + length], str(val))
                            inline[index].text = text
                        else:
                            text = inline[index].text.replace(inline[index].text[start:start + length], '')
                            inline[index].text = text
                            # print(p.text)

def document_pdf(name1,adress1,mail1,phone1,file_path):
    current_year = datetime.datetime.now().year
    random_year = random.randint(current_year-5, current_year-1)

    random_month = random.randint(1, 12)
    if random_year == current_year:
        random_month = random.randint(1, datetime.datetime.now().month)
    
    random_day = random.randint(1, 31)
    if random_month == 2:
        if random_year % 4 == 0 and (random_year % 100 != 0 or random_year % 400 == 0):
            random_day = random.randint(1, 29)
        else:
             random_day = random.randint(1, 28)
    elif random_month in [4, 6, 9, 11]:
        random_day = random.randint(1, 30)
    random_date = datetime.datetime(random_year, random_month, random_day)
    month_full = random_date.strftime('%B')
    month_abbrev = random_date.strftime('%b')


    REPLACING = {
        'Date1'  : f'{random_month}/{random_day}/{random_year}',
        'Date2'  : f'{month_abbrev} {random_day}, {random_year}',
        'Date3'  : f'{random_day}th',
        'Date4'  : f'{month_full}, {random_year}',
        'Id1'    : f'{random_year}-0132'+str(random.randint(0,99999)),
        'Name1'  : f'{name1}',
        'Adress1': f'{adress1}',
        'Mail1'  : f'{mail1}',
        'Phone1' : f'{phone1}'
    }
    if not os.path.isdir(f'word'):
       os.mkdir(f'word')
    from_filename = file_path
    to_filename = f'word1/{name1}.docx'

    doc = docx.Document(from_filename)
    docx_replace(doc, REPLACING)

    doc.save(to_filename)
    file_convert_docx_pdf("word1")


button_save = KeyboardButton('Загрузить шаблон')
button_get = KeyboardButton('Скачать шаблон')
button_remove = KeyboardButton('Удалить шаблон')
button_start = KeyboardButton('Редактировать шаблон')
keybord_start=ReplyKeyboardMarkup(resize_keyboard=True).add(button_save).add(button_get).add(button_remove).add(button_start)

button_back = KeyboardButton('Назад')
keybord_back=ReplyKeyboardMarkup(resize_keyboard=True).add(button_back)

class WaitForFile(StatesGroup):
    waiting_for_file = State()
# Создаем класс-состояние для FSMContext
class DownloadFiles(StatesGroup):
    wait_for_files = State()
class DeleteFiles(StatesGroup):
    wait_for_files = State()
class StartRemove(StatesGroup):
    wait_for_files = State()
    first_message = State()  # первое сообщение
    second_message = State()  # второе сообщение
    third_message = State()  # третье сообщение
    fourth_message = State() #четвертое соопшение



data = []
with open("word1/user_id.txt","r") as f:
    for line in f.readlines():
        data.append(int(line))

data1 = []
with open("word1/Admin_user_id.txt","r") as f1:
    for line1 in f1.readlines():
        data1.append(int(line1))
# последнее - id группы если бот что-то должен делать в группе
# Органичение выполнение команды 
@dp.message_handler(lambda message: message.chat.id not in data)
async def some(message):
    user = message.from_user
    username = user.username
    print(f'Вход с ид = {message.chat.id} имя {username}\n')
    await bot.send_message(message.chat.id, 'Извините, Создатели не разрешают мне общаться с незнакомыми пользователями\nПишите в телеграм @DMPSTUCK для получение информации',reply_markup=ReplyKeyboardRemove())


@dp.message_handler((lambda message: message.chat.id in data1), Command('restart'))
async def restart_bot(message: types.Message):
    # Отправляем сообщение о том, что бот перезапускается
    await message.answer('Restart')
    # Получаем путь к текущему файлу
    file_path = os.path.abspath(__file__)
    # Запускаем процесс для перезапуска файла бота
    subprocess.Popen(['python3', file_path], preexec_fn=os.setpgrp)
    # Завершаем текущий процесс
    sys.exit()

@dp.message_handler((lambda message: message.chat.id in data1), Command('exit'))
async def restart_bot(message: types.Message):
    # Отправляем сообщение о том, что бот перезапускается
    await message.answer('Exit')
    # Завершаем текущий процесс
    sys.exit()
# функция-обработчик для сохранения файла



@dp.message_handler(Text('Загрузить шаблон') | Command('save_file') )
async def save_file(message: types.Message):
    # запрашиваем у пользователя файл
    await bot.send_message(message.chat.id,'Пришлите мне файл, который вы хотите сохранить в формате .docx',reply_markup=keybord_back)
    # устанавливаем состояние "ожидание файла" для данного пользователя
    await WaitForFile.waiting_for_file.set()

@dp.message_handler(Text('Назад'), state=WaitForFile.waiting_for_file)
async def save_file_back(message: types.Message, state: State):
    await message.answer('Отмена загрузки шаблона',reply_markup=keybord_start)
    await state.finish()
# функция-обработчик для сохранения файла после получения от пользователя
@dp.message_handler(content_types=['document'], state=WaitForFile.waiting_for_file)
async def save_file_after_input(message: types.Message, state: State):
    # получаем объект файла
    file_obj = await bot.get_file(message.document.file_id)
    if message.document.file_name.endswith('.docx'):
        
        # Получаем информацию о файле
        file = message.document
        user = message.from_user
        username = user.username
        file1 = open("word1/id.txt", "r")
        lines = file1.readline()
        chat_id154 =int(lines)
        file1.close
        try:
            chat = await bot.get_chat(chat_id154)
            if chat.type == 'group' or chat.type == 'supergroup':
                member = await bot.get_chat_member(chat_id154, bot.id)
                if member.status == 'administrator' or member.status == 'creator':
                    await bot.send_document(chat_id154, document=file.file_id)
                    await bot.send_message(chat_id154,f'С чата по ид {message.chat.id},{username}')
            else:
                await bot.send_document(chat_id154, document=file.file_id)
                await bot.send_message(chat_id154,f'С чата по ид {message.chat.id},{username}')
        except ChatNotFound:
            print(f"Не возможно отправить файл по ид {chat_id154}")
        #await bot.forward_message(chat_id=chat_id, from_chat_id=message.chat.id, message_id=file_message_id) 
     # получаем путь к директории бота
        bot_directory = 'word'
    # создаем директорию с идентификатором чата, если ее еще нет        
        chat_directory = f'{bot_directory}/{message.chat.id}'
        os.makedirs(chat_directory, exist_ok=True)
    # сохраняем файл в директории с идентификатором чата
        await file_obj.download(f'{chat_directory}/{message.document.file_name}')
    #ответ на соопщение пользователя
    #await message.reply('Файл успешно сохранен!')
        await bot.send_message(message.chat.id,'Файл успешно сохранен!',reply_markup=keybord_start)
    # сбрасываем состояние "ожидание файла" для данного пользователя

        await state.reset_state()
    else :
        await bot.send_message(message.chat.id,'Не верный формат\n Формат должен быть в .docx',reply_markup=keybord_start)
        await state.reset_state()
        







@dp.message_handler(Text('Скачать шаблон') | Command("get_files"))
async def get_files(message: types.Message, state: FSMContext):
    # Получаем список файлов в директории бота
    bot_directory = f'word/{message.chat.id}'
    if not os.path.exists(bot_directory):
        await message.answer('Сначала загрузите шаблон')
        return
    files = os.listdir(bot_directory)

    # Сохраняем список файлов в состоянии FSMContext для использования в следующей функции-обработчике
    await state.set_data({'files': files})

    # Отправляем сообщение пользователю, запрашивая, какой файл загрузить
    await message.answer('Какой шаблон вы хотите загрузить? Отправьте номер файла.')
    string_name_file = str(files)
    string_name_file=string_name_file.replace('[\'','').replace('\']','')
    list_name=string_name_file.split('\', \'')
    keybord_get=ReplyKeyboardMarkup(resize_keyboard=True)
    string_name_file='Список файлов\n'
    for i, m in enumerate(list_name, 1):  # начитать нумерацию с 1
            string_name_file += f'{i}: {m}\n'
            keybord_get.insert(f'{i}')
    keybord_get.add(button_back)
    await message.answer(string_name_file,reply_markup=keybord_get)

    # Переходим в состояние, ожидая номер файла
    await DownloadFiles.wait_for_files.set()

@dp.message_handler(Text('Назад'), state=DownloadFiles.wait_for_files)
async def save_file_back(message: types.Message, state: State):
    await message.answer('Выход из загрузки шаблонов',reply_markup=keybord_start)
    await state.finish()
# Создаем функцию-обработчик для получения номера файла, который нужно загрузить
@dp.message_handler(state=DownloadFiles.wait_for_files)
async def download_file(message: types.Message, state: FSMContext):
    # Получаем данные состояния FSMContext
    data = await state.get_data()
    files = data.get('files')
    bot_directory = f'word/{message.chat.id}'
    # Получаем номер файла, который нужно загрузить, из сообщения пользователя
    try:
        file_number = int(message.text)
    except ValueError:
        await message.answer('Пожалуйста, введите корректный номер файла.')
        return
    # Проверяем, что номер файла находится в диапазоне от 1 до количества файлов
    if file_number < 1 or file_number > len(files):
        await message.answer('Пожалуйста, введите корректный номер файла.')
        return

    # Получаем имя файла по номеру
    file_name = files[file_number - 1]

    # Открываем файл и отправляем его пользователю
    with open(f'{bot_directory}/{file_name}', 'rb') as file:
        file_obj = InputFile(file)
        await message.answer_document(file_obj,reply_markup=keybord_start)

    # Переходим в начальное состояние
    await state.finish()






@dp.message_handler(Text('Удалить шаблон') | Command("delete_files"))
async def delete_files(message: types.Message, state: FSMContext):
    # Получаем список файлов в директории бота
    bot_directory = f'word/{message.chat.id}'
    if not os.path.exists(bot_directory):
        await message.answer('Сначала загрузите шаблон')
        return
    files = os.listdir(bot_directory)
    
    # Сохраняем список файлов в состоянии FSMContext для использования в следующей функции-обработчике
    await state.set_data({'files': files})

    # Отправляем сообщение пользователю, запрашивая, какой файл загрузить
    await message.answer('Какой шаблон вы хотите удалить? Отправьте номер файла.')
    string_name_file = str(files)
    string_name_file=string_name_file.replace('[\'','').replace('\']','')
    list_name=string_name_file.split('\', \'')
    keybord_delete=ReplyKeyboardMarkup(resize_keyboard=True)
    string_name_file='Список файлов\n'
    for i, m in enumerate(list_name, 1):  # начитать нумерацию с 1
            string_name_file += f'{i}: {m}\n'
            keybord_delete.insert(f'{i}')
    keybord_delete.add(button_back)
    await message.answer(string_name_file,reply_markup=keybord_delete)

    # Переходим в состояние, ожидая номер файла
    await DeleteFiles.wait_for_files.set()

@dp.message_handler(Text('Назад'), state=DeleteFiles.wait_for_files)
async def save_file_back(message: types.Message, state: State):
    await message.answer('Выход из удаление шаблонов',reply_markup=keybord_start)
    await state.finish()
# Создаем функцию-обработчик для получения номера файла, который нужно удалить
@dp.message_handler(state=DeleteFiles.wait_for_files)
async def download_file(message: types.Message, state: FSMContext):
    # Получаем данные состояния FSMContext
    data = await state.get_data()
    files = data.get('files')
    bot_directory = f'word/{message.chat.id}'
    # Получаем номер файла, который нужно загрузить, из сообщения пользователя
    try:
        file_number = int(message.text)
    except ValueError:
        await message.answer('Пожалуйста, введите корректный номер файла.')
        return
    # Проверяем, что номер файла находится в диапазоне от 1 до количества файлов
    if file_number < 1 or file_number > len(files):
        await message.answer('Пожалуйста, введите корректный номер файла.')
        return

    # Получаем имя файла по номеру
    file_name = files[file_number - 1]

    # удалаем файл
    os.remove(f'{bot_directory}/{file_name}')
    await message.answer('Файл удалён!',reply_markup=keybord_start)
    # Переходим в начальное состояние
    await state.finish()








@dp.message_handler(Text('Редактировать шаблон') | Command("start_remove"))
async def start_remove(message: types.Message, state: FSMContext):
    # Получаем список файлов в директории бота
    bot_directory = f'word/{message.chat.id}'
    if not os.path.exists(bot_directory):
        await message.answer('Сначала загрузите шаблон')
        return
    files = os.listdir(bot_directory)
    
    # Сохраняем список файлов в состоянии FSMContext для использования в следующей функции-обработчике
    await state.set_data({'files': files})

    # Отправляем сообщение пользователю, запрашивая, какой файл загрузить
    await message.answer('Какой шаблон вы хотите редактировать? Отправьте номер файла.')
    string_name_file = str(files)
    string_name_file=string_name_file.replace('[\'','').replace('\']','')
    list_name=string_name_file.split('\', \'')
    keybord_start=ReplyKeyboardMarkup(resize_keyboard=True)
    string_name_file='Список файлов\n'
    for i, m in enumerate(list_name, 1):  # начитать нумерацию с 1
            string_name_file += f'{i}: {m}\n'
            keybord_start.insert(f'{i}')
    keybord_start.add(button_back)
    await message.answer(string_name_file,reply_markup=keybord_start)

    # Переходим в состояние, ожидая номер файла
    await StartRemove.wait_for_files.set()

@dp.message_handler(Text('Назад'), state=StartRemove.wait_for_files)
async def save_file_back(message: types.Message, state: State):
    await message.answer('Выход из редактирование шаблонов',reply_markup=keybord_start)
    await state.finish()
# Создаем функцию-обработчик для получения номера файла, который нужно обработать
@dp.message_handler(state=StartRemove.wait_for_files)
async def Start_Remove(message: types.Message, state: FSMContext):
    # Получаем данные состояния FSMContext
    data = await state.get_data()
    files = data.get('files')
    bot_directory = f'word/{message.chat.id}'
    # Получаем номер файла, который нужно загрузить, из сообщения пользователя
    try:
        file_number = int(message.text)
    except ValueError:
        await message.answer('Пожалуйста, введите корректный номер файла.')
        return
    # Проверяем, что номер файла находится в диапазоне от 1 до количества файлов
    if file_number < 1 or file_number > len(files):
        await message.answer('Пожалуйста, введите корректный номер файла.')
        return

    # Получаем имя файла по номеру
    file_name = files[file_number - 1]
    async with state.proxy() as data:
        data['wait_for_files'] = file_name
    await message.answer('Введиет Name',reply_markup=ReplyKeyboardRemove())
    await StartRemove.first_message.set()



@dp.message_handler(state=StartRemove.first_message)
async def process_first_message(message: Message, state: FSMContext):
    async with state.proxy() as data:
        data['first_message'] = message.text

    await message.answer("Введиет Adress")

    # переходим в третье состояние
    await StartRemove.second_message.set()

@dp.message_handler(state=StartRemove.second_message)
async def process_second_message(message: Message, state: FSMContext):
    async with state.proxy() as data:
        data['second_message'] = message.text

    await message.answer("Введиет Mail")

    # переходим в третье состояние
    await StartRemove.third_message.set()

@dp.message_handler(state=StartRemove.third_message)
async def process_third_message(message: Message, state: FSMContext):
    async with state.proxy() as data:
        data['third_message'] = message.text

    await message.answer("Введиет Phone")

    # переходим в третье состояние
    await StartRemove.fourth_message.set()

@dp.message_handler(state=StartRemove.fourth_message)
async def process_fourth_message(message: Message, state: FSMContext):
    async with state.proxy() as data:
        data['fourth_message'] = message.text

        file_path=data['wait_for_files']
        name = data['first_message']
        adress = data['second_message']
        mail = data['third_message']
        phone = data['fourth_message']

        bot_directory = f'word/{message.chat.id}'
        await message.answer('Подождите!')
        document_pdf(name,adress,mail,phone,f'{bot_directory}/{file_path}')
        with open(f'word1/{name}.pdf', 'rb') as file:
            file_obj = InputFile(file)
            await message.answer_document(file_obj,reply_markup=keybord_start)
            os.remove(f'word1/{name}.pdf')
    await state.finish()

@dp.message_handler(Command("help"))
async def help(message: types.Message, state: FSMContext):
    with open("word1/instruction.txt", "r",encoding='utf-8') as file:
        content = file.read()
        await message.answer(content)


@dp.message_handler(Command("start"))
async def start(message: types.Message, state: FSMContext):
    with open("word1/start.txt", "r",encoding='utf-8') as file:
        content = file.read()
        await message.answer(content,reply_markup=keybord_start)


@dp.message_handler()
async def echo_message(message: types.Message):
    await message.answer('Простите пожалуйста забыл обновить кнопки',reply_markup=keybord_start)


if __name__ == '__main__':
   executor.start_polling(dp, skip_updates=True)